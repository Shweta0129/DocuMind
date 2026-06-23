"""DOCX export for DocuMind AI."""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import base64

_DATA_IMG_RE = re.compile(r"^\s*!\[[^\]]*\]\(\s*(data:image/[A-Za-z0-9.+-]+;base64,([A-Za-z0-9+/=\s]+))\)\s*$")


def _add_table_from_md(doc, md_table_lines):
    rows = [
        [c.strip() for c in line.strip().strip("|").split("|")]
        for line in md_table_lines
        if line.strip()
    ]
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:]  # skip separator
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, r in enumerate(body):
        for j in range(len(header)):
            table.rows[i + 1].cells[j].text = r[j] if j < len(r) else ""


def _add_inline(p, text):
    """Inline markdown rendering (bold/italic/code) into a paragraph."""
    # Bold first
    tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for tk in tokens:
        if not tk:
            continue
        if tk.startswith("**") and tk.endswith("**"):
            r = p.add_run(tk[2:-2])
            r.bold = True
        elif tk.startswith("`") and tk.endswith("`"):
            r = p.add_run(tk[1:-1])
            r.font.name = "Consolas"
        else:
            p.add_run(tk)


def _render_markdown(doc, md):
    if not md:
        return
    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # fenced code block (```mermaid ... ```): we can't render Mermaid to an
        # image server-side, so emit a labeled monospaced block. Diagrams render
        # fully in the in-app viewer; user images below DO embed as pictures.
        fence = re.match(r"^\s*```\s*([\w-]*)\s*$", line)
        if fence:
            lang = (fence.group(1) or "").lower()
            i += 1
            code = []
            while i < len(lines) and not re.match(r"^\s*```\s*$", lines[i]):
                code.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            label = "Diagram (Mermaid — view interactive version in the app)" if lang == "mermaid" else (lang or "code")
            lp = doc.add_paragraph()
            lr = lp.add_run(label)
            lr.italic = True
            lr.font.size = Pt(9)
            lr.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)
            for cl in code:
                cp = doc.add_paragraph()
                cr = cp.add_run(cl)
                cr.font.name = "Consolas"
                cr.font.size = Pt(9)
            continue
        # embedded user image (data-URI) -> real picture
        img = _DATA_IMG_RE.match(line)
        if img:
            try:
                raw = base64.b64decode(re.sub(r"\s+", "", img.group(2)))
                doc.add_picture(BytesIO(raw), width=Inches(5.8))
            except Exception:
                pass
            i += 1
            continue
        # table — a row is any line containing "|"; the 2nd line must be a separator.
        # Rows may omit leading/trailing pipes (models often do).
        def _is_sep(s):
            return "-" in s and re.match(r"^\s*\|?[\s:|-]+\|?\s*$", s) is not None
        if "|" in line and i + 1 < len(lines) and _is_sep(lines[i + 1]):
            tbl = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip() and not re.match(r"^\s*#{1,6}\s", lines[i]):
                tbl.append(lines[i])
                i += 1
            _add_table_from_md(doc, tbl)
            continue
        # heading
        h = re.match(r"^(#{1,6})\s+(.+)$", line)
        if h:
            lvl = min(len(h.group(1)), 4)
            doc.add_heading(h.group(2), level=lvl)
            i += 1
            continue
        # list
        if re.match(r"^\s*[-*]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                p = doc.add_paragraph(style="List Number")
                _add_inline(p, re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            continue
        # blockquote
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _add_inline(p, line.lstrip(">").strip())
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        # paragraph (collect)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6}\s|>|\s*[-*]\s|\s*\d+\.\s|\s*\|)", lines[i]):
            para_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        _add_inline(p, " ".join(para_lines))


def _add_page_number(paragraph):
    """Add a Word PAGE field to a paragraph (for footer page numbers)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_control_block(doc, document, s):
    """Document Control: company/version properties + author/reviewer/approver sign-off."""
    from datetime import datetime, timezone

    doc.add_heading("Document Control", level=2)
    props = []
    if s.get("company_name"):
        props.append(("Company", s["company_name"]))
    if s.get("project_name"):
        props.append(("Project", s["project_name"]))
    if s.get("document_id"):
        props.append(("Document ID", s["document_id"]))
    props.append(("Document Type", (document.get("type") or "").upper()))
    props.append(("Version", f"v{s.get('version_number') or document.get('version_number', '1.0')}"))
    props.append(("Date", datetime.now(timezone.utc).strftime("%Y-%m-%d")))

    pt = doc.add_table(rows=0, cols=2)
    pt.style = "Light Grid Accent 1"
    for k, v in props:
        cells = pt.add_row().cells
        cells[0].text = k
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
        cells[1].text = str(v)

    signers = [("Author", s.get("author")), ("Reviewer", s.get("reviewer")), ("Approver", s.get("approver"))]
    if any(name for _, name in signers):
        doc.add_paragraph()
        st = doc.add_table(rows=1, cols=4)
        st.style = "Light Grid Accent 1"
        for j, h in enumerate(["Role", "Name", "Signature", "Date"]):
            st.rows[0].cells[j].text = h
            for run in st.rows[0].cells[j].paragraphs[0].runs:
                run.bold = True
        for role, name in signers:
            cells = st.add_row().cells
            cells[0].text = role
            cells[1].text = name or ""
    doc.add_paragraph().add_run("─" * 60)


def build_docx(document, settings=None) -> bytes:
    """Build a .docx file from a stored Document object and optional header/footer settings."""
    s = settings or {}
    doc = Document()
    # ---- Styles ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Header ----
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_bits = []
    if s.get("company_name"):
        header_bits.append(s["company_name"])
    if s.get("project_name"):
        header_bits.append(s["project_name"])
    if s.get("document_id"):
        header_bits.append(f"Doc ID: {s['document_id']}")
    hp.text = "  ·  ".join(header_bits) if header_bits else (s.get("company_name") or "DocuMind AI")

    # ---- Footer ----
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_bits = []
    if s.get("version_number"):
        footer_bits.append(f"v{s['version_number']}")
    if document.get("version_number"):
        footer_bits.append(f"v{document['version_number']}")
    if s.get("author"):
        footer_bits.append(f"Author: {s['author']}")
    if s.get("reviewer"):
        footer_bits.append(f"Reviewer: {s['reviewer']}")
    if s.get("approver"):
        footer_bits.append(f"Approver: {s['approver']}")
    fp.add_run("  ·  ".join(footer_bits) if footer_bits else "Generated by DocuMind AI")
    fp.add_run("    Page ")
    _add_page_number(fp)

    # ---- Title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run(document.get("title") or "Untitled Document")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x0F, 0x0F, 0x0F)

    # Meta line
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"{(document.get('type') or '').upper()}   ·   "
        f"Quality {document.get('completeness_score', '—')}%   ·   "
        f"v{document.get('version_number', '1.0')}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x6E)

    doc.add_paragraph().add_run("─" * 60)

    # Document Control (company header/footer info + sign-off)
    _add_control_block(doc, document, s)

    # Sections
    for sec in document.get("content", {}).get("sections", []):
        doc.add_heading(sec.get("heading", "Section"), level=1)
        _render_markdown(doc, sec.get("content", ""))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ---------------- File text extraction (for reviewer & templates) ----------------
def extract_text_from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(data))
    out = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(out)


def extract_text_from_docx(data: bytes) -> str:
    f = BytesIO(data)
    doc = Document(f)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)
